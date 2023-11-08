from __future__ import annotations # for Python3.9 type hinting compatibility

import pathlib as pl

import geopandas as gpd
import docx
import PIL
from PIL import Image

ROOT = pl.Path(__file__).parent

def build_report(
    data: gpd.GeoDataFrame,
    long_response_fields: list[str] = None,
    image_fields: list[str] = None,
    image_dir: str | pl.Path = "",
    portrait_image_size: dict[str, float] = {"height": 12},
    landscape_image_size: dict[str, float] = {"width": 9},
    intro_text: str = None,
    intro_image_path: str = None,
    template_path: str = None,
    *,
    dropna: bool = True,
    use_page_breaks: bool = True,
    page_break_after_intro: bool = True,
    use_two_column_tables: bool = False,
) -> docx.document.Document:
    """
    This function builds a Microsoft Word report containing all of the features in 
    `data`. The document is formatted as an 'index' of features. For each feature, there
    is: a table of its attributes (excluding the `long_response_fields`), followed by
    sentences of data (from the `long_response_fields`), followed by images of that 
    feature saved to the `image_dir` folder and with filenames defined by `image_fields`.

    Inputs:
    - `data`: a GeoPandas GeoDataFrame to detail in the report
    - `long_response_fields`: a list of fields from `data` to include as 'normal' text 
    below the table instead of in the table
    - `image_fields`: a list of fields from `data` that contain filenames for the images
    to be included
    - `image_dir`: directory where the images are stored
    - `portrait_image_size`: maximum size (in cm) for portrait images in the document, 
    can provide height or width but not both (default = `{"height": 12}`)
    - `landscape_image_size`: maximum size (in cm) for landscape images in the document, 
    can provide height or width but not both (default = `{"width": 9}`)
    - `intro_text`: some introductory text to include at the start of the document
    - `intro_image_path`: path to an image to include after the introductory text
    - `template_path`: path to an empty Microsoft Word document to use as a template for
    fonts, styles etc.
    - `dropna`: whether to drop NA values in the document (default = `True`)
    - `use_page_breaks`: whether to include page breaks between features (default = `True`)
    - `page_break_after_intro`: whether to include a page break after the intro text
    - `use_two_column_tables`: whether to duplicate the [field: value] columns horizontally
    (default = `False`)
    """

    # Check inputs
    if not long_response_fields:
        long_response_fields = []
    if not image_fields:
        image_fields = []
    if type(image_dir) == str:
        image_dir = pl.Path(image_dir)

    # Initialise document
    doc = docx.Document(template_path)

    # Introduction
    intro = False
    if intro_text:
        doc.add_paragraph(intro_text)
        intro = True
    if intro_image_path:
        doc.add_picture(str(intro_image_path))
        intro = True
    if intro and page_break_after_intro:
        doc.add_page_break()

    # Loop through features and add them to the report
    for i, iterrow in data.iterrows():

        # Define table contents
        row = (
            iterrow
            .drop(long_response_fields)
            .drop(image_fields)
            .drop("geometry")
        )
        if dropna:
            row = row.dropna()

        # Build the table
        if len(row) > 0:
            # Set the column sizes
            if use_two_column_tables and len(row) > 3:
                column_groups = [0, 2]
                columns = [
                    {"name": "Field", "width": 5},
                    {"name": "Value", "width": 3.3},
                    {"name": "Field", "width": 5},
                    {"name": "Value", "width": 3.3},
                ]
            else:
                column_groups = [0]
                columns = [
                    {"name": "Field", "width": 5},
                    {"name": "Value", "width": 11},
                ]

            # Initialise table
            table = doc.add_table(rows=1, cols=len(columns)) 
            table.allow_autofit = False
            hdr_cells = table.rows[0].cells
            for j in range(len(columns)):
                hdr_cells[j].text = columns[j]["name"]
                table.columns[j].width = docx.shared.Cm(columns[j]["width"])

            # Populate table
            j = 0
            while j < len(row):
                table_row = table.add_row().cells
                for k in column_groups:
                    if j < len(row):
                        table_row[k+0].text = row.index[j]
                        table_row[k+1].text = str(row.iloc[j])
                    j += 1

        # Add long response fields outside the table
        doc.add_paragraph()
        for f in long_response_fields:
            if iterrow[f]:
                paragraph = doc.add_paragraph()
                paragraph.add_run([f, ": "]).bold = True
                paragraph.add_run(iterrow[f])

        # Add images
        for f in image_fields:
            if f in iterrow and iterrow[f] is not None:
                doc.add_paragraph()
                paragraph = doc.add_paragraph()
                # Check image orientation
                path = str(image_dir / iterrow[f])
                img = PIL.Image.open(path)
                w, h = img.size
                args = {}
                if h > w:
                    # Portrait
                    if "height" in portrait_image_size:
                        args["height"] = docx.shared.Cm(portrait_image_size["height"])
                    if "width" in portrait_image_size:
                        args["width"] = docx.shared.Cm(portrait_image_size["width"])
                else:
                    # Landscape
                    if "height" in landscape_image_size:
                        args["height"] = docx.shared.Cm(landscape_image_size["height"])
                    if "width" in landscape_image_size:
                        args["width"] = docx.shared.Cm(landscape_image_size["width"])
        
                # Add to document
                run = paragraph.add_run()
                run.add_picture(path, **args)

        if use_page_breaks and i < len(data) - 1:
            doc.add_page_break()
        else:
            doc.add_paragraph()

    return doc
